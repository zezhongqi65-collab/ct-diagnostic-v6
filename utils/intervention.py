"""
干预出题模块 — 根据诊断结果生成个性化思维训练题单（纸质版）。

流程：诊断结果（优先干预维度）→ 按维度分组学生 → 选题/生成变式 → Word 题单。

题目来源（混合模式）：
- 预置题库 data/question_bank.json 打底（保证质量稳定）；
- 可选调用 DeepSeek 为母题生成变式（换数字/情境，扩充题量）；
- DeepSeek 不可用时自动回退到母题，保证服务不中断。

参考：清华 OpenMAIC 的「根据学生特点个性化出题」思路，
结合本系统「无电脑、纸质作答、基础薄弱」的约束，采用不插电思维训练题。
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path
from typing import Optional

_PARENT = Path(__file__).resolve().parent.parent
if str(_PARENT) not in sys.path:
    sys.path.insert(0, str(_PARENT))

# ── 常量（与 completeV6_patched 保持一致）────────────────
CAUSAL_ORDER: list[str] = ['抽象', '分解', '算法设计', '建模', '评估']
DIM_NAMES: dict[str, str] = {
    '抽象': '模式识别与抽象',
    '分解': '问题分解',
    '算法设计': '算法与步骤设计',
    '建模': '数学建模与符号化',
    '评估': '方案评估与调试',
}


def _resource_path(relative: str) -> Path:
    """资源读取路径：兼容源码运行与 PyInstaller 打包（sys._MEIPASS）。"""
    base = Path(getattr(sys, '_MEIPASS', Path(__file__).resolve().parent.parent))
    return base / relative


def _writable_dir(relative: str) -> Path:
    """可写目录：打包后用 exe 所在目录，源码运行用项目根目录。"""
    if getattr(sys, 'frozen', False):
        base = Path(sys.executable).resolve().parent
    else:
        base = Path(__file__).resolve().parent.parent
    return base / relative


QUESTION_BANK_PATH: Path = _resource_path('data/question_bank.json')
HISTORY_PATH: Path = _writable_dir('data/intervention_history.json')


def load_history() -> dict:
    """加载出题历史（{周次: {维度: [已用题目id]}}）。"""
    if HISTORY_PATH.exists():
        try:
            with open(HISTORY_PATH, encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_history(history: dict) -> None:
    """保存出题历史。"""
    HISTORY_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(HISTORY_PATH, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def used_ids(history: dict, dim: str) -> set[str]:
    """某维度在所有周次里已经用过的题目 id 集合。"""
    used: set[str] = set()
    for week_dims in history.values():
        for qid in week_dims.get(dim, []):
            used.add(qid)
    return used


def reset_history() -> None:
    """清空出题历史（新学期重新开始一轮）。"""
    save_history({})


# ═══════════════════════════════════════════════════════
# 1. 题库加载与诊断结果解析
# ═══════════════════════════════════════════════════════


def load_question_bank(path: Optional[Path] = None) -> dict[str, list[dict]]:
    """加载题库 JSON，返回 {维度: [题目dict, ...]}。"""
    bank_path = path or QUESTION_BANK_PATH
    with open(bank_path, encoding='utf-8') as f:
        return json.load(f)


def extract_priority_dims(diag_result: dict) -> list[str]:
    """从单个诊断结果提取「优先干预」维度（✅分类）。

    df_results 已按优先级排序，因此多个优先维度时按顺序返回。
    """
    df = diag_result.get('df_results')
    if df is None or len(df) == 0:
        return []
    priority = df[df['分类'].str.startswith('✅')]
    return priority['维度'].tolist()


def group_students_by_dim(diagnosis_results: list[dict]) -> dict[str, list[str]]:
    """按优先干预维度分组学生。

    每个学生只归入「最优先的一个维度」（df_results 已排序）。
    无优先维度的学生归入「综合」组（做混合基础题）。

    Returns:
        {维度: [student_id, ...]}，只包含有学生的维度。
    """
    groups: dict[str, list[str]] = {dim: [] for dim in CAUSAL_ORDER}
    groups['综合'] = []

    for r in diagnosis_results:
        dims = extract_priority_dims(r)
        if dims:
            groups[dims[0]].append(str(r['student_id']))
        else:
            groups['综合'].append(str(r['student_id']))

    return {dim: students for dim, students in groups.items() if students}


# ═══════════════════════════════════════════════════════
# 2. 题目变式生成（DeepSeek，可选）
# ═══════════════════════════════════════════════════════


def _deepseek_generate(system: str, user: str) -> Optional[str]:
    """调用 DeepSeek 生成文本，失败返回 None。"""
    api_key = os.environ.get('DEEPSEEK_API_KEY')
    if not api_key:
        return None
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url='https://api.deepseek.com')
        response = client.chat.completions.create(
            model='deepseek-chat',
            messages=[
                {'role': 'system', 'content': system},
                {'role': 'user', 'content': user},
            ],
            temperature=0.7,
            timeout=30,
        )
        return response.choices[0].message.content
    except Exception:
        return None


def generate_question_variant(question: dict) -> Optional[dict]:
    """为一道母题生成一道同类型、同难度的变式题。

    Returns:
        变式题 dict（含 stem/answer/hint/is_variant=True），失败返回 None。
    """
    system = (
        '你是一名高中数学思维训练教师，擅长出简单有趣、适合纸笔作答的思维题。'
        '你只输出题目，不输出无关内容。'
    )
    user = (
        f'请根据下面这道题，出一道同类型、同难度的变式题（只换数字或情境，'
        f'题型和考察点保持不变）。\n\n'
        f'原题题型：{question["type"]}\n'
        f'原题题干：{question["stem"]}\n'
        f'变式方向：{question.get("variant", "自由发挥")}\n\n'
        f'请按以下格式输出（不要用 Markdown 代码块）：\n'
        f'题干：...\n'
        f'答案：...\n'
        f'提示：...'
    )
    text = _deepseek_generate(system, user)
    if not text:
        return None

    stem = answer = hint = None
    for line in text.splitlines():
        line = line.strip()
        if line.startswith('题干') and '：' in line:
            stem = line.split('：', 1)[1].strip()
        elif line.startswith('答案') and '：' in line:
            answer = line.split('：', 1)[1].strip()
        elif line.startswith('提示') and '：' in line:
            hint = line.split('：', 1)[1].strip()
    if not stem:
        return None

    return {
        'id': question.get('id', '') + '-v',
        'type': question['type'],
        'stem': stem,
        'answer': answer or '（请教师核对）',
        'hint': hint or '',
        'difficulty': question.get('difficulty', 2),
        'is_variant': True,
    }


def select_questions(dim: str, count: int, use_llm: bool,
                     bank: dict[str, list[dict]],
                     used: set[str] | None = None) -> list[dict]:
    """为某维度（或「综合」）选出 count 道题。

    - 「综合」组（无优先维度的学生）：从五个维度混合抽取基础题。
    - use_llm=True 时，尝试为母题生成变式（成功用变式，失败回退母题）。
    - used：已发过的题目 id 集合，自动跳过（跨周不重复）；题库不足时循环补足。
    """
    if dim == '综合':
        pool = [q for d in CAUSAL_ORDER for q in bank.get(d, [])]
    else:
        pool = bank.get(dim, [])
    if not pool:
        return []

    if used:
        fresh = [q for q in pool if q.get('id', '') not in used]
        if len(fresh) >= count:
            pool = fresh  # 未用题目足够，只取未用
        else:
            # 未用优先，用完的放后面兜底
            pool = fresh + [q for q in pool if q.get('id', '') in used]

    if use_llm:
        selected: list[dict] = []
        for q in pool[:count]:
            variant = generate_question_variant(q)
            selected.append(variant if variant else q)
        return selected[:count]

    return [pool[i % len(pool)] for i in range(count)]


# ═══════════════════════════════════════════════════════
# 3. Word 题单生成（python-docx）
# ═══════════════════════════════════════════════════════


def _set_cjk_font(doc) -> None:
    """设置 Word 默认字体为中文字体（宋体）。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _add_heading(doc, text: str) -> None:
    doc.add_heading(text, level=1)


def _add_question(doc, idx: int, q: dict) -> None:
    """向文档追加一道题（题干 + 作答区）。"""
    tag = '（AI 生成，请教师核对）' if q.get('is_variant') else ''
    doc.add_paragraph(f'第 {idx} 题　【{q["type"]}】{tag}')
    # 题干按换行符拆成多行（九宫格、步骤排序等题有 \n）
    for line in q['stem'].split('\n'):
        doc.add_paragraph(line)
    # 作答空白区
    for _ in range(3):
        doc.add_paragraph('')
    doc.add_paragraph('')


def _add_answer_section(doc, questions: list[dict]) -> None:
    """追加参考答案页（教师用）。"""
    doc.add_page_break()
    _add_heading(doc, '参考答案（教师用）')
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f'第 {i} 题：{q["answer"]}')
        if q.get('hint'):
            doc.add_paragraph(f'　　提示：{q["hint"]}')


def build_dimension_sheet(dim: str, student_ids: list[str],
                          questions: list[dict], week: int):
    """生成某维度的 Word 题单，返回 Document 对象。"""
    from docx import Document

    doc = Document()
    _set_cjk_font(doc)

    full_name = DIM_NAMES.get(dim, dim)
    _add_heading(doc, f'{full_name} 思维训练题单（第 {week} 周）')

    doc.add_paragraph(f'训练维度：{full_name}')
    doc.add_paragraph(f'学生名单（{len(student_ids)} 人）：' + '、'.join(student_ids))
    doc.add_paragraph('作答要求：请独立思考，用文字或图形写出你的思路。')
    doc.add_paragraph('')

    for i, q in enumerate(questions, 1):
        _add_question(doc, i, q)

    _add_answer_section(doc, questions)
    return doc


# ═══════════════════════════════════════════════════════
# 4. 主入口
# ═══════════════════════════════════════════════════════


def generate_intervention_package(
    diagnosis_results: list[dict],
    week: int,
    use_llm: bool = False,
    questions_per_sheet: int = 5,
    out_dir: Optional[Path] = None,
    bank: Optional[dict[str, list[dict]]] = None,
) -> list[Path]:
    """根据诊断结果生成所有维度的 Word 题单。

    Args:
        diagnosis_results: 批量诊断结果列表（每项含 student_id 与 df_results）。
        week: 第几周（用于题单标题）。
        use_llm: 是否用 DeepSeek 生成变式。
        questions_per_sheet: 每份题单的题数。
        out_dir: 输出目录，默认 data/intervention/。
        bank: 题库（默认从 data/question_bank.json 加载）。

    Returns:
        生成的 Word 文件路径列表。
    """
    bank = bank if bank is not None else load_question_bank()
    history = load_history()
    week_key = str(week)
    used_now = history.setdefault(week_key, {})
    groups = group_students_by_dim(diagnosis_results)

    out_dir = out_dir or _writable_dir('data/intervention')
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    paths: list[Path] = []
    for dim, student_ids in groups.items():
        used = used_ids(history, dim)
        questions = select_questions(dim, questions_per_sheet, use_llm, bank, used=used)
        if not questions:
            continue
        used_now[dim] = list(dict.fromkeys(used_now.get(dim, []) + [q.get('id', '') for q in questions]))
        doc = build_dimension_sheet(dim, student_ids, questions, week)
        safe_dim = dim if dim == '综合' else dim
        path = out_dir / f'第{week}周_{safe_dim}训练题单.docx'
        doc.save(path)
        paths.append(path)

    save_history(history)
    return paths


# ═══════════════════════════════════════════════════════
# 5. 一人一单（按学生生成个性化题单，合并成一份 Word）
# ═══════════════════════════════════════════════════════


def generate_student_packets(
    diagnosis_results: list[dict],
    week: int,
    questions_per_dim: int = 10,
    max_dims: int = 2,
    out_dir: Optional[Path] = None,
    bank: Optional[dict[str, list[dict]]] = None,
) -> tuple[Path, list[dict]]:
    """按学生生成「一人一单」：合并成一份 Word（每名学生一页）+ 分发清单。

    每个学生取「优先干预」维度中排序最靠前的 max_dims 个，各 questions_per_dim 题。
    同维度学生共享本周题目；跨周自动跳过已发过的题。

    Returns:
        (合并后的 Word 题单路径, 分发清单行列表 [{学生ID, 训练维度, 题数}]).
    """
    bank = bank if bank is not None else load_question_bank()
    history = load_history()
    week_key = str(week)
    used_now = history.setdefault(week_key, {})

    # 1) 每个学生的薄弱维度（优先干预，按优先级）
    student_dims: dict[str, list[str]] = {}
    for r in diagnosis_results:
        dims = extract_priority_dims(r)
        student_dims[str(r['student_id'])] = dims[:max_dims]

    # 2) 涉及的所有维度各选一次题（同维度学生共享本周题目）
    all_dims = sorted({d for dims in student_dims.values() for d in dims})
    dim_questions: dict[str, list[dict]] = {}
    for dim in all_dims:
        used = used_ids(history, dim)
        qs = select_questions(dim, questions_per_dim, False, bank, used=used)
        used_now[dim] = list(dict.fromkeys(used_now.get(dim, []) + [q.get('id', '') for q in qs]))
        dim_questions[dim] = qs
    save_history(history)

    # 3) 合并成一份 Word
    from docx import Document
    doc = Document()
    _set_cjk_font(doc)
    _add_heading(doc, f'计算思维个性化训练题单（第 {week} 周）—— 一人一单')

    dist_rows: list[dict] = []
    first = True
    for sid, dims in student_dims.items():
        if not dims:
            continue
        if not first:
            doc.add_page_break()
        first = False

        doc.add_heading(str(sid), level=1)
        doc.add_paragraph('训练维度：' + '、'.join(dims))
        qs_all = [(d, q) for d in dims for q in dim_questions.get(d, [])]
        idx = 1
        for d in dims:
            doc.add_heading(DIM_NAMES.get(d, d), level=2)
            for _, q in [x for x in qs_all if x[0] == d]:
                _add_question(doc, idx, q)
                idx += 1
        dist_rows.append({'学生ID': str(sid), '训练维度': '、'.join(dims), '题数': idx - 1})

    # 4) 参考答案页（教师用，同维度题目按 id 去重）
    doc.add_page_break()
    _add_heading(doc, '参考答案（教师用）')
    seen: set[str] = set()
    ai = 1
    for dim, qs in dim_questions.items():
        for q in qs:
            if q.get('id') in seen:
                continue
            seen.add(q.get('id', ''))
            doc.add_paragraph(f'第 {ai} 题（{dim}·{q["type"]}）：{q["answer"]}')
            if q.get('hint'):
                doc.add_paragraph(f'　　提示：{q["hint"]}')
            ai += 1

    out_dir = Path(out_dir) if out_dir else _writable_dir('data/intervention')
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f'第{week}周_一人一单题单.docx'
    doc.save(path)
    return path, dist_rows
